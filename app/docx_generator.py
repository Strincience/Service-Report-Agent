from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generate_docx(data: dict, output_path: str):
    """
    Generate a DOCX file from JSON schema data.
    
    Args:
        data: Dictionary containing the report schema with heading, workers_charge, service_outline, 
              sermon, attendance, challenges, and authorship information.
        output_path: Path where the DOCX file will be saved.
    """
    document = Document()

    # Base font
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # Title
    title = document.add_paragraph("CHAPEL OF PEACE AND JOY")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(14)
    
    subtitle = document.add_paragraph("LEAD CITY UNIVERSITY")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True
    subtitle.runs[0].font.size = Pt(12)
    
    report_title = document.add_paragraph("Sunday Service Report")
    report_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    report_title.runs[0].bold = True
    report_title.runs[0].font.size = Pt(12)
    
    document.add_paragraph("")  # Spacing

    # HEADING SECTION
    if data.get("heading"):
        heading = data["heading"]
        p = document.add_paragraph("Heading")
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(12)
        
        if heading.get("service_date"):
            document.add_paragraph(f"Service Date: {heading['service_date']}", style="List Bullet")
        if heading.get("service_theme"):
            document.add_paragraph(f"Service Theme: {heading['service_theme']}", style="List Bullet")
        if heading.get("start_time"):
            document.add_paragraph(f"Start Time: {heading['start_time']}", style="List Bullet")
        if heading.get("venue"):
            document.add_paragraph(f"Venue: {heading['venue']}", style="List Bullet")
        
        document.add_paragraph("")  # Spacing

    # WORKERS' CHARGE PROGRAMME OUTLINE
    if data.get("workers_charge"):
        workers = data["workers_charge"]
        p = document.add_paragraph("Section A – Workers' Charge Programme Outline")
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(11)
        
        if workers.get("opening_time"):
            document.add_paragraph(f"Opening Time: {workers['opening_time']}", style="List Bullet")
        if workers.get("opening_prayer_and_worship"):
            document.add_paragraph(f"Opening Prayer & Worship: {workers['opening_prayer_and_worship']}", style="List Bullet")
        if workers.get("sessions_and_facilitators"):
            document.add_paragraph(f"Sessions & Facilitators: {workers['sessions_and_facilitators']}", style="List Bullet")
        if workers.get("key_highlights"):
            document.add_paragraph(f"Key Highlights: {workers['key_highlights']}", style="List Bullet")
        
        document.add_paragraph("")  # Spacing

    # SERVICE OUTLINE
    if data.get("service_outline"):
        service = data["service_outline"]
        p = document.add_paragraph("Section B – Service Outline")
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(11)
        
        if service.get("order_of_service"):
            document.add_paragraph(service["order_of_service"])
        
        document.add_paragraph("")  # Spacing

    # SERMON HIGHLIGHTS
    if data.get("sermon"):
        sermon = data["sermon"]
        p = document.add_paragraph("Section C – Sermon Highlights")
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(11)
        
        if sermon.get("sermon_topic") or sermon.get("topic"):
            topic = sermon.get("sermon_topic") or sermon.get("topic")
            document.add_paragraph(f"Sermon Topic: {topic}", style="List Bullet")
        if sermon.get("preacher"):
            document.add_paragraph(f"Preacher: {sermon['preacher']}", style="List Bullet")
        if sermon.get("scriptures"):
            document.add_paragraph(f"Scripture(s): {sermon['scriptures']}", style="List Bullet")
        if sermon.get("main_thought"):
            document.add_paragraph(f"Main Thought: {sermon['main_thought']}", style="List Bullet")
        if sermon.get("key_points"):
            document.add_paragraph(f"Key Points: {sermon['key_points']}", style="List Bullet")
        
        document.add_paragraph("")  # Spacing

    # ATTENDANCE
    if data.get("attendance"):
        attendance = data["attendance"]
        p = document.add_paragraph("Section D – Attendance")
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(11)
        
        male = attendance.get("male", 0)
        female = attendance.get("female", 0)
        total = attendance.get("total", 0)
        
        document.add_paragraph(f"Male: {male}", style="List Bullet")
        document.add_paragraph(f"Female: {female}", style="List Bullet")
        document.add_paragraph(f"Total: {total}", style="List Bullet")
        
        document.add_paragraph("")  # Spacing

    # CHALLENGES
    if data.get("challenges"):
        challenges = data["challenges"]
        p = document.add_paragraph("Section E – Challenges Encountered")
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(11)
        
        if challenges.get("had_challenges"):
            document.add_paragraph(f"Were There Any Challenges?: {challenges['had_challenges']}", style="List Bullet")
        if challenges.get("description"):
            document.add_paragraph(f"Description: {challenges['description']}", style="List Bullet")
        
        document.add_paragraph("")  # Spacing

    # AUTHORSHIP
    if data.get("authorship"):
        auth = data["authorship"]
        p = document.add_paragraph("Section F – Session Authorship")
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(11)
        
        if auth.get("prepared_by"):
            document.add_paragraph(f"Report Prepared By: {auth['prepared_by']}", style="List Bullet")
        if auth.get("role"):
            document.add_paragraph(f"Role / Unit: {auth['role']}", style="List Bullet")
        if auth.get("timestamp"):
            document.add_paragraph(f"Date of Submission: {auth['timestamp']}", style="List Bullet")

    document.save(output_path)
    